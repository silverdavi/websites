#!/usr/bin/env python3
"""
Convert Markdown files to PDF and DOCX for USPTO filing.

Generates:
- SPECIFICATION.pdf / SPECIFICATION.docx  (from draft.md)
- CLAIMS.pdf / CLAIMS.docx (from CLAIMS.md)
- COVER_SHEET.pdf (from COVER_SHEET.md)

Usage:
    python md_to_pdf.py

Requirements:
    pip install markdown weasyprint python-docx
"""

import markdown
from weasyprint import HTML, CSS
from pathlib import Path

# Try to import docx for Word conversion
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("Note: python-docx not installed. DOCX output disabled.")
    print("      Install with: pip install python-docx")

CSS_STYLE = """
@page {
    size: letter;
    margin: 1in;
}
body {
    font-family: "Times New Roman", Times, serif;
    font-size: 12pt;
    line-height: 1.5;
}
h1 {
    font-size: 16pt;
    font-weight: bold;
    margin-top: 24pt;
    margin-bottom: 12pt;
    text-align: center;
}
h2 {
    font-size: 14pt;
    font-weight: bold;
    margin-top: 18pt;
    margin-bottom: 10pt;
}
h3 {
    font-size: 12pt;
    font-weight: bold;
    margin-top: 14pt;
    margin-bottom: 8pt;
}
h4 {
    font-size: 11pt;
    font-weight: bold;
    margin-top: 12pt;
    margin-bottom: 6pt;
}
p {
    margin: 8pt 0;
    text-align: justify;
}
table {
    border-collapse: collapse;
    margin: 10pt 0;
    width: 100%;
}
th, td {
    border: 1px solid black;
    padding: 6pt 8pt;
    text-align: left;
}
th {
    background-color: #f0f0f0;
}
code {
    font-family: "Courier New", Courier, monospace;
    font-size: 10pt;
    background-color: #f5f5f5;
    padding: 2pt 4pt;
}
pre {
    font-family: "Courier New", Courier, monospace;
    font-size: 10pt;
    background-color: #f5f5f5;
    padding: 10pt;
    margin: 10pt 0;
    white-space: pre-wrap;
    word-wrap: break-word;
}
blockquote {
    border-left: 3pt solid #ccc;
    margin-left: 0;
    padding-left: 12pt;
    font-style: italic;
}
hr {
    border: none;
    border-top: 1pt solid #ccc;
    margin: 20pt 0;
}
ul, ol {
    margin: 10pt 0;
    padding-left: 20pt;
}
li {
    margin: 4pt 0;
}
strong {
    font-weight: bold;
}
em {
    font-style: italic;
}
"""


def md_to_pdf(input_file: str, output_file: str):
    """Convert a markdown file to PDF."""
    md_content = Path(input_file).read_text()
    
    # Convert markdown to HTML
    html_content = markdown.markdown(
        md_content,
        extensions=['tables', 'fenced_code', 'toc', 'nl2br']
    )
    
    # Wrap in HTML document
    full_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
    </head>
    <body>
    {html_content}
    </body>
    </html>
    """
    
    # Convert to PDF
    HTML(string=full_html).write_pdf(
        output_file,
        stylesheets=[CSS(string=CSS_STYLE)]
    )
    print(f"Created: {output_file}")


def md_to_docx(input_file: str, output_file: str):
    """Convert a markdown file to DOCX (simple conversion)."""
    if not HAS_DOCX:
        print(f"Skipping DOCX: {output_file} (python-docx not installed)")
        return
    
    md_content = Path(input_file).read_text()
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Process markdown line by line (simple parser)
    lines = md_content.split('\n')
    in_code_block = False
    
    for line in lines:
        stripped = line.strip()
        
        # Code blocks
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            continue
        
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'Normal'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            continue
        
        # Headers
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:].replace('**', ''), level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:].replace('**', ''), level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:].replace('**', ''), level=3)
        elif stripped.startswith('#### '):
            doc.add_heading(stripped[5:].replace('**', ''), level=4)
        # Horizontal rule
        elif stripped == '---':
            doc.add_paragraph('_' * 50)
        # Empty line
        elif stripped == '':
            continue
        # Tables (basic handling)
        elif stripped.startswith('|'):
            # Just add as text for now
            p = doc.add_paragraph(stripped)
        # List items
        elif stripped.startswith('- ') or stripped.startswith('* '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] == '.':
            doc.add_paragraph(stripped[3:], style='List Number')
        # Regular paragraph
        else:
            # Clean up markdown formatting
            text = stripped.replace('**', '').replace('*', '').replace('`', '')
            if text:
                doc.add_paragraph(text)
    
    doc.save(output_file)
    print(f"Created: {output_file}")


if __name__ == '__main__':
    print("=" * 60)
    print("Converting Markdown to PDF/DOCX for USPTO Filing")
    print("=" * 60)
    
    # Convert specification
    if Path('draft.md').exists():
        md_to_pdf('draft.md', 'SPECIFICATION.pdf')
        md_to_docx('draft.md', 'SPECIFICATION.docx')
    
    # Convert claims
    if Path('CLAIMS.md').exists():
        md_to_pdf('CLAIMS.md', 'CLAIMS.pdf')
        md_to_docx('CLAIMS.md', 'CLAIMS.docx')
    
    # Convert cover sheet
    if Path('COVER_SHEET.md').exists():
        md_to_pdf('COVER_SHEET.md', 'COVER_SHEET.pdf')
    
    # Convert filing cover sheet
    if Path('COVER_SHEET_FILING.md').exists():
        md_to_pdf('COVER_SHEET_FILING.md', 'COVER_SHEET_FILING.pdf')
    
    # Convert prior art report
    if Path('PRIOR_ART_REPORT.md').exists():
        md_to_pdf('PRIOR_ART_REPORT.md', 'PRIOR_ART_REPORT.pdf')
    
    # Convert IDS references list
    if Path('IDS_REFERENCES.md').exists():
        md_to_pdf('IDS_REFERENCES.md', 'IDS_REFERENCES.pdf')
    
    print("=" * 60)
    print("USPTO Filing Documents Ready:")
    print("=" * 60)
    print("  1. SPECIFICATION.pdf / SPECIFICATION.docx")
    print("  2. CLAIMS.pdf / CLAIMS.docx")
    print("  3. COVER_SHEET.pdf")
    print("  4. DRAWINGS.pdf (combine fig*.pdf)")
    print("  5. PRIOR_ART_REPORT.pdf (internal reference)")
    print("  6. IDS_REFERENCES.pdf (for IDS filing)")
    print("")
    print("Upload all to USPTO Patent Center.")
    print("=" * 60)
