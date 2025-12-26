#!/usr/bin/env python3
"""
Convert Markdown files to USPTO-compliant DOCX and PDF.
Includes paragraph numbering, page numbers, proper claim format.

Run: python convert_docs.py
"""

import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("Error: python-docx not installed. Run: pip install python-docx")

try:
    import markdown
    from weasyprint import HTML, CSS
    HAS_WEASYPRINT = True
except ImportError:
    HAS_WEASYPRINT = False


def add_page_numbers(doc):
    """Add page numbers to document footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        
        # Create paragraph for page number
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add "Page X of Y" field
        run = p.add_run()
        
        # PAGE field
        fld_char1 = OxmlElement('w:fldChar')
        fld_char1.set(qn('w:fldCharType'), 'begin')
        
        instr_text = OxmlElement('w:instrText')
        instr_text.text = "PAGE"
        
        fld_char2 = OxmlElement('w:fldChar')
        fld_char2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fld_char1)
        run._r.append(instr_text)
        run._r.append(fld_char2)


def md_to_docx_uspto(input_file: str, output_file: str):
    """Convert markdown to USPTO-compliant DOCX."""
    if not HAS_DOCX:
        print(f"  Skipping DOCX: {output_file} (python-docx not installed)")
        return
    
    md_content = Path(input_file).read_text()
    doc = Document()
    
    # Set up page margins (1 inch all around for USPTO)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    
    # Set up Normal style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    
    # Process content
    lines = md_content.split('\n')
    in_code_block = False
    code_buffer = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Skip horizontal rules
        if line.strip() == '---':
            i += 1
            continue
        
        # Code blocks
        if line.startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_buffer)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue
        
        # Main title (# PROVISIONAL...)
        if line.startswith('# '):
            text = line[2:].strip().replace('**', '')
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
        
        # Section headers (## FIELD OF...)
        elif line.startswith('## '):
            text = line[3:].strip().replace('**', '')
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(12)
        
        # Subsection headers (### System Architecture)
        elif line.startswith('### '):
            text = line[4:].strip().replace('**', '')
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(12)
        
        # Bold field labels (**Title:**)
        elif line.startswith('**') and ':**' in line:
            text = line.replace('**', '')
            p = doc.add_paragraph()
            # Split on first colon
            if ':' in text:
                parts = text.split(':', 1)
                run = p.add_run(parts[0] + ':')
                run.bold = True
                if len(parts) > 1:
                    p.add_run(parts[1])
        
        # Paragraph numbers [0001]
        elif line.strip().startswith('[0'):
            # Extract paragraph number and text
            match = re.match(r'\[(\d+)\]\s*(.*)', line.strip())
            if match:
                para_num = match.group(1)
                text = match.group(2)
                p = doc.add_paragraph()
                run = p.add_run(f'[{para_num}] ')
                run.bold = False
                # Clean up text
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
                p.add_run(text)
                p.paragraph_format.first_line_indent = Inches(0.5)
        
        # Claims (numbered lines starting with digit and period)
        elif re.match(r'^\d+\.\s*\(', line.strip()):
            # Claim line like "1. (Original) An automated..."
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            
            # Parse claim number, status, and text
            match = re.match(r'^(\d+)\.\s*\((\w+)\)\s*(.*)', line.strip())
            if match:
                claim_num = match.group(1)
                status = match.group(2)
                claim_text = match.group(3)
                
                run = p.add_run(f'{claim_num}. ({status}) ')
                run.bold = False
                p.add_run(claim_text)
        
        # Numbered list items (1. A natural-language...)
        elif re.match(r'^\d+\.\s', line.strip()) and not line.strip().startswith('['):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Number')
        
        # Bullet points
        elif line.strip().startswith('* ') or line.strip().startswith('- '):
            text = line.strip()[2:]
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            doc.add_paragraph(text, style='List Bullet')
        
        # Regular paragraphs
        elif line.strip() and not line.startswith('#'):
            text = line.strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            if text and text != '*End of Specification*':
                p = doc.add_paragraph(text)
                p.paragraph_format.first_line_indent = Inches(0.5)
        
        i += 1
    
    # Add page numbers
    add_page_numbers(doc)
    
    doc.save(output_file)
    print(f"  Created: {output_file}")


def md_to_pdf(input_file: str, output_file: str):
    """Convert markdown to PDF."""
    if not HAS_WEASYPRINT:
        print(f"  Skipping PDF: {output_file} (weasyprint not installed)")
        return
    
    md_content = Path(input_file).read_text()
    
    # Convert markdown to HTML
    html_content = markdown.markdown(
        md_content,
        extensions=['tables', 'fenced_code', 'toc']
    )
    
    css_style = """
    @page {
        size: letter;
        margin: 1in;
        @bottom-center {
            content: counter(page);
        }
    }
    body {
        font-family: "Times New Roman", Times, serif;
        font-size: 12pt;
        line-height: 1.5;
    }
    h1 { font-size: 14pt; font-weight: bold; text-align: center; }
    h2 { font-size: 12pt; font-weight: bold; text-align: center; margin-top: 18pt; }
    h3 { font-size: 12pt; font-weight: bold; margin-top: 12pt; }
    p { text-indent: 0.5in; }
    code { font-family: "Courier New", monospace; font-size: 10pt; }
    """
    
    full_html = f"""
    <!DOCTYPE html>
    <html>
    <head><meta charset="utf-8"></head>
    <body>{html_content}</body>
    </html>
    """
    
    HTML(string=full_html).write_pdf(output_file, stylesheets=[CSS(string=css_style)])
    print(f"  Created: {output_file}")


if __name__ == '__main__':
    print("Converting USPTO filing documents...")
    print("=" * 50)
    
    print("\n1. SPECIFICATION:")
    md_to_docx_uspto('provisional.md', 'SPECIFICATION.docx')
    md_to_pdf('provisional.md', 'SPECIFICATION.pdf')
    
    print("\n2. COVER SHEET:")
    md_to_docx_uspto('cover_sheet.md', 'COVER_SHEET.docx')
    md_to_pdf('cover_sheet.md', 'COVER_SHEET.pdf')
    
    print("\n" + "=" * 50)
    print("USPTO Filing Documents Ready:")
    print("  - SPECIFICATION.docx (with paragraph numbers & page numbers)")
    print("  - COVER_SHEET.docx")
    print("  - DRAWINGS.pdf")
    print("\nUpload to: https://patentcenter.uspto.gov")
